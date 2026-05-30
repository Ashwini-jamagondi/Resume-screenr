import os
import re
from typing import Optional, List, Dict, Any
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_doc(file_path: str) -> str:
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except Exception:
        pass
    try:
        from docx import Document
        return extract_text_from_docx(file_path)
    except Exception as e:
        logger.error(f"DOC extraction error: {e}")
        return ""

def extract_text(file_path: str, file_type: str) -> str:
    file_type = file_type.lower().lstrip(".")
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type == "doc":
        return extract_text_from_doc(file_path)
    return ""

def extract_email(text: str) -> Optional[str]:
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None

def extract_phone(text: str) -> Optional[str]:
    pattern = r'(\+?[\d\s\-\(\)]{10,15})'
    matches = re.findall(pattern, text)
    phones = [m.strip() for m in matches if len(re.sub(r'\D', '', m)) >= 10]
    return phones[0] if phones else None

def extract_name(text: str) -> str:
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:
        if (len(line.split()) <= 5 and
            not any(char.isdigit() for char in line) and
            not '@' in line and
            len(line) > 2 and
            line.replace(' ', '').replace('-', '').isalpha()):
            return line
    return lines[0] if lines else "Unknown Candidate"

TECH_SKILLS = [
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
    "react", "vue", "angular", "node.js", "express", "django", "flask",
    "fastapi", "spring", "sql", "mysql", "postgresql", "mongodb", "redis",
    "docker", "kubernetes", "aws", "azure", "gcp", "git", "linux",
    "machine learning", "deep learning", "tensorflow", "pytorch", "scikit-learn",
    "data analysis", "pandas", "numpy", "tableau", "power bi",
    "html", "css", "rest api", "graphql", "microservices", "agile", "scrum",
    "devops", "ci/cd", "jenkins", "terraform", "ansible"
]

def extract_skills(text: str) -> List[str]:
    text_lower = text.lower()
    found_skills = []
    for skill in TECH_SKILLS:
        if skill in text_lower:
            found_skills.append(skill)
    return list(set(found_skills))

def extract_experience_years(text: str) -> Optional[float]:
    patterns = [
        r'(\d+)\+?\s*years?\s+of\s+experience',
        r'(\d+)\+?\s*years?\s+experience',
        r'experience\s+of\s+(\d+)\+?\s*years?',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text.lower())
        if matches:
            return float(matches[0])
    return None

def extract_education(text: str) -> Optional[str]:
    degrees = ["phd", "ph.d", "doctorate", "master", "m.s.", "m.tech", "mba",
               "bachelor", "b.s.", "b.tech", "b.e.", "b.sc", "associate", "diploma"]
    text_lower = text.lower()
    for degree in degrees:
        if degree in text_lower:
            return degree.upper().replace(".", "")
    return None

def parse_resume(file_path: str, file_type: str) -> Dict[str, Any]:
    text = extract_text(file_path, file_type)
    return {
        "extracted_text": text,
        "candidate_name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": extract_skills(text),
        "experience_years": extract_experience_years(text),
        "education": extract_education(text),
    }
